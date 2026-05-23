"""
Generate IOC Management Documentation as a DOCX file.
Run: python generate_ioc_docs.py
Output: IOC_Management_Documentation.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_heading(paragraph, text, level=1):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    if not paragraph.runs:
        run = paragraph.runs[0]
    else:
        run.text = text
    if level == 1:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0E, 0x74, 0xD4)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x37, 0x51, 0xA0)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0E, 0x74, 0xD4)
        run.font.size = Pt(22)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(15)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x37, 0x51, 0xA0)
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1A56DB')
        tcPr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = 'F0F4FF' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tcPr.append(shd)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A56DB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)

# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SecureEye Portal")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0E, 0x74, 0xD4)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("IOC Management Module")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Technical Documentation")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "Table of Contents")
toc_items = [
    ("1", "Overview of IOC Management"),
    ("2", "Tab 1 — Live Raw Feed"),
    ("2.1", "What is the Raw Feed?"),
    ("2.2", "Raw Feed Sources (14 Sources)"),
    ("2.3", "How Raw Feed Works"),
    ("2.4", "Filtering & Pagination"),
    ("3", "Tab 2 — Enriched IOC + Risk Score"),
    ("3.1", "What is Enrichment?"),
    ("3.2", "Enrichment Sources"),
    ("3.3", "Risk Score Calculation"),
    ("3.4", "VirusTotal-Style Confirmation"),
    ("3.5", "False Positive Removal"),
    ("3.6", "Rate Limiting"),
    ("3.7", "Step-by-Step Enrichment Flow"),
    ("4", "Tab 3 — Live Search"),
    ("5", "Tab 4 — Tracked Indicators"),
    ("6", "Backend Architecture"),
    ("7", "API Endpoints"),
    ("8", "Environment Variables & API Keys"),
    ("9", "Limitations & Notes"),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"  {num}   {item}")
    run.font.size = Pt(11)
    if '.' not in num:
        run.font.bold = True
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Overview of IOC Management")
divider(doc)

body(doc,
    "The IOC Management module is the central threat intelligence hub of the SecureEye portal. "
    "It provides security analysts with real-time access to global threat indicators, enriched "
    "intelligence from multiple sources, a live search capability, and a local database of "
    "tracked indicators specific to the organization."
)

body(doc,
    "An Indicator of Compromise (IOC) is a piece of forensic evidence that suggests a network "
    "or system may have been breached. Common IOC types include:"
)
bullet(doc, "IP Addresses — Malicious servers, C2 (command and control) endpoints, attacker IPs")
bullet(doc, "Domains — Malicious domains used for phishing, malware distribution, C2 communication")
bullet(doc, "File Hashes (MD5/SHA256) — Fingerprints of known malware samples")
bullet(doc, "URLs — Specific web addresses hosting malware, phishing pages, or exploit kits")

doc.add_paragraph()
h3(doc, "Module Structure — 4 Tabs")
add_table(doc,
    ["Tab", "Name", "Purpose", "Data Source"],
    [
        ["Tab 1", "Live Raw Feed",           "Browse 1000s of real-time IOCs from 14 global feeds", "External threat intel feeds (auto-fetched)"],
        ["Tab 2", "Enriched + Risk Score",   "Confirm malicious IOCs with multi-source enrichment",  "Shodan, AbuseIPDB, GreyNoise, MalwareBazaar, OTX"],
        ["Tab 3", "Live Search",             "Search any IOC across local DB + external sources",    "Local DB + OTX + Threat intelligence APIs"],
        ["Tab 4", "Tracked Indicators",      "Organization-specific IOC watchlist",                   "Local PostgreSQL database"],
    ],
    col_widths=[0.6, 1.5, 2.8, 2.3]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — RAW FEED
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Tab 1 — Live Raw Feed")
divider(doc)

h2(doc, "2.1 What is the Raw Feed?")
body(doc,
    "The Live Raw Feed aggregates IOCs in real-time from 14 external Open Source Intelligence "
    "(OSINT) and threat intelligence feeds. These are global, community-maintained databases "
    "that continuously publish known malicious indicators. The feed is fetched fresh every time "
    "the tab is opened or the refresh button is clicked."
)
body(doc,
    "Raw IOCs are not filtered by whether they affect your organization specifically — they represent "
    "the global threat landscape. Think of it as a continuously updated global blocklist."
)

h2(doc, "2.2 Raw Feed Sources (14 Sources)")

h3(doc, "Group 1 — abuse.ch Suite (6 Sources)")
body(doc,
    "abuse.ch is a Swiss non-profit project that tracks cybercriminal infrastructure. "
    "All sources are 100% free with no API key required."
)
add_table(doc,
    ["Source", "IOC Type", "What it Tracks", "Update Frequency"],
    [
        ["URLHaus",       "URLs",    "Malware delivery URLs, phishing links",               "Every 5 minutes"],
        ["FeodoTracker",  "IPs",     "Botnet C2 IPs (Emotet, QakBot, TrickBot, Dridex)",   "Every 5 minutes"],
        ["FeodoDomains",  "Domains", "Botnet C2 domains from the same botnets",             "Every 5 minutes"],
        ["MalwareBazaar", "Hashes",  "SHA256 hashes of recent malware samples",             "Continuous"],
        ["ThreatFox",     "Mixed",   "IPs, domains, URLs, hashes from threat actors",       "Continuous"],
        ["SSL Blacklist",  "Domains", "Malicious SSL certificate domains, botnet C2",        "Every 30 minutes"],
    ],
    col_widths=[1.3, 0.9, 2.8, 1.6]
)

h3(doc, "Group 2 — IP Intelligence (6 Sources)")
add_table(doc,
    ["Source", "IOC Type", "What it Tracks", "Notes"],
    [
        ["DShield / SANS",    "IPs", "Top attacking IPs from global honeypot network",    "Curated daily top-1000 attacker list"],
        ["Blocklist.de",      "IPs", "fail2ban-reported attacker IPs worldwide",          "Community-sourced, very large list"],
        ["Spamhaus DROP",     "IPs", "Hijacked IP space used for spam and attacks",       "Network CIDR ranges — filtered from enrichment"],
        ["C2 Tracker",        "IPs", "Active C2 servers tracked by the community",        "GitHub-maintained, daily updated"],
        ["Emerging Threats",  "IPs", "Proofpoint compromised IP reputation list",         "Industry-grade threat intelligence"],
        ["Tor Exit Nodes",    "IPs", "All current Tor exit relay IP addresses",           "Updated hourly from Tor Project"],
    ],
    col_widths=[1.3, 0.9, 2.8, 1.6]
)

h3(doc, "Group 3 — Phishing (2 Sources)")
add_table(doc,
    ["Source", "IOC Type", "What it Tracks", "Notes"],
    [
        ["OpenPhish",  "URLs", "Community phishing URLs, brand impersonation",     "No API key required"],
        ["PhishTank",  "URLs", "Verified phishing URLs with brand & target data",  "Community-verified, high accuracy"],
    ],
    col_widths=[1.3, 0.9, 2.8, 1.6]
)

h2(doc, "2.3 How Raw Feed Works")
body(doc, "The backend fetches all 14 sources simultaneously using Python asyncio (async/await):")
code_block(doc, "fetch_all_raw_iocs()  →  asyncio.gather(*all_14_sources)  →  normalize  →  merge  →  return")
body(doc, "Each source fetcher:")
bullet(doc, "Makes an HTTP request to the feed URL")
bullet(doc, "Parses the response (CSV, JSON, or plain text)")
bullet(doc, "Normalizes each IOC into a standard dictionary:")
code_block(doc, '{ "value": "1.2.3.4", "ioc_type": "ip", "severity": "high", "feed": "FeodoTracker", "threat": "Emotet", "tags": ["botnet", "c2"] }')
bullet(doc, "Handles errors gracefully — if one source fails, others still return data")

h2(doc, "2.4 Filtering & Pagination")
body(doc, "Users can filter the raw feed by:")
bullet(doc, "IOC Type — IP / Domain / Hash / URL")
bullet(doc, "Severity — Critical / High / Medium / Low")
bullet(doc, "Source — Select specific feed from grouped dropdown")
body(doc, "The feed displays 25 IOCs per page with Previous/Next navigation. Total IOCs from all sources can exceed 5,000.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ENRICHED IOC
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Tab 2 — Enriched IOC + Risk Score")
divider(doc)

h2(doc, "3.1 What is Enrichment?")
body(doc,
    "Enrichment takes raw IOCs (which are just values like '1.2.3.4') and queries multiple "
    "threat intelligence databases to answer the question: 'Is this IOC actually malicious, "
    "and how dangerous is it?'"
)
body(doc,
    "This is conceptually similar to VirusTotal — where VT submits a file to 70+ antivirus engines, "
    "our enrichment system submits an IOC to multiple reputation databases and produces a "
    "composite Risk Score from 0 to 100."
)

add_table(doc,
    ["VirusTotal", "SecureEye Enrichment"],
    [
        ["Scans with 70+ AV engines",             "Queries 4-5 reputation databases"],
        ["Output: '34/72 engines detected'",       "Output: '2/4 sources detected'"],
        ["Risk verdict: MALICIOUS / CLEAN",        "Risk Score: 0-100 + CRITICAL/HIGH/MEDIUM/LOW/SAFE"],
        ["Premium: $200-$500/month",               "100% Free"],
        ["Rate limit: 500 lookups/day (free tier)","Rate limit: 200/type/hour (self-imposed)"],
    ],
    col_widths=[3.5, 3.5]
)

h2(doc, "3.2 Enrichment Sources")
body(doc, "The following sources are queried for each IOC type:")

h3(doc, "For IP Addresses (4 Sources)")
add_table(doc,
    ["Source", "API Key", "Daily Limit", "What it Returns"],
    [
        ["Shodan InternetDB", "None (free)",  "Unlimited",  "Open ports, CVEs on the IP, tags (c2/compromised/honeypot)"],
        ["AbuseIPDB",         "Free key",     "1,000/day",  "Abuse confidence score 0-100%, number of abuse reports, ISP"],
        ["GreyNoise Community","None (free)", "Unlimited",  "Classification: malicious / benign / unknown"],
        ["ip-api.com",        "None (free)",  "45 req/min", "Country, ISP, proxy detection, hosting provider detection"],
    ],
    col_widths=[1.5, 1.0, 1.0, 3.2]
)

h3(doc, "For File Hashes (2 Sources)")
add_table(doc,
    ["Source", "API Key", "Daily Limit", "What it Returns"],
    [
        ["MalwareBazaar",     "None (free)", "Unlimited", "Malware family name, file type, malware tags"],
        ["AlienVault OTX",    "Free key",    "Unlimited", "Number of threat pulses mentioning this hash"],
    ],
    col_widths=[1.5, 1.0, 1.0, 3.2]
)

h3(doc, "For Domains & URLs (1 Source each)")
add_table(doc,
    ["IOC Type", "Source", "API Key", "What it Returns"],
    [
        ["Domain", "AlienVault OTX", "Free key", "Threat pulse count, malware campaigns, threat actors using this domain"],
        ["URL",    "AlienVault OTX", "Free key", "Threat pulses linked to this specific URL"],
    ],
    col_widths=[1.0, 1.5, 1.0, 3.2]
)

h2(doc, "3.3 Risk Score Calculation")
body(doc, "The risk score is calculated as:")
code_block(doc, "FINAL_SCORE = BASE_SCORE + sum(DELTA from each source)  [clamped to 0-100]")

h3(doc, "Base Score (from feed severity)")
add_table(doc,
    ["Feed Severity", "Base Score"],
    [
        ["Critical", "70"],
        ["High",     "55"],
        ["Medium",   "35"],
        ["Low",      "15"],
    ],
    col_widths=[2.0, 1.5]
)

h3(doc, "Delta Values (per source finding)")
add_table(doc,
    ["Source", "Condition", "Score Delta"],
    [
        ["Shodan",     "Each CVE found on IP",              "+8 per CVE (max +30)"],
        ["Shodan",     "Tag = 'c2' or 'compromised'",       "+25"],
        ["Shodan",     "Tag = 'self-signed' or 'honeypot'", "+10"],
        ["Shodan",     "Suspicious open ports (22,445,etc)","+(5 per port)"],
        ["AbuseIPDB",  "Abuse confidence score N%",         "+(N × 0.8)"],
        ["GreyNoise",  "Classification = malicious",        "+35"],
        ["GreyNoise",  "Classification = benign",           "-20 (score reduced!)"],
        ["ip-api",     "is_proxy = true",                   "+15"],
        ["ip-api",     "is_hosting = true",                 "+8"],
        ["MalwareBazaar","Hash found in database",          "+75"],
        ["OTX",        "pulse_count = N",                   "+(N × 6, max +40)"],
    ],
    col_widths=[1.4, 3.0, 1.8]
)

h3(doc, "Risk Labels")
add_table(doc,
    ["Score Range", "Label", "Color"],
    [
        ["85 — 100", "CRITICAL", "Red"],
        ["65 — 84",  "HIGH",     "Orange"],
        ["40 — 64",  "MEDIUM",   "Yellow"],
        ["15 — 39",  "LOW",      "Green"],
        ["0  — 14",  "SAFE",     "Cyan"],
    ],
    col_widths=[1.5, 1.5, 1.5]
)

h2(doc, "3.4 VirusTotal-Style Confirmation")
body(doc,
    "Like VirusTotal's 'X/N engines detected' display, the Enriched tab shows a detection ratio "
    "for each IOC:"
)
code_block(doc, "detection_ratio = f'{confirmation_count}/{len(sources_checked)} sources'")
body(doc, "Example: '3/4 sources detected' means 3 out of 4 queried sources confirmed this IOC as malicious.")

h2(doc, "3.5 False Positive Removal")
body(doc,
    "This is the key difference from the Raw Feed. The Enriched tab ONLY shows IOCs that "
    "pass all three conditions:"
)
bullet(doc, "confirmed_malicious = True  →  at least 1 enrichment source flagged it")
bullet(doc, "risk_score ≥ 30  →  minimum meaningful risk level")
bullet(doc, "Not a CIDR block  →  Spamhaus DROP gives network ranges like 5.188.0.0/23, not individual IPs")
body(doc, "IOCs that fail any of these conditions are counted as 'false positives removed' and shown in the header.")

h2(doc, "3.6 Rate Limiting")
body(doc,
    "To be respectful to the free APIs and avoid bans, the system limits enrichment to "
    "200 IOCs per type per hour. This is a self-imposed limit and can be changed in the configuration."
)
add_table(doc,
    ["IOC Type", "Max per Hour", "Max per Day"],
    [
        ["IP",     "200", "4,800"],
        ["Domain", "200", "4,800"],
        ["Hash",   "200", "4,800"],
        ["URL",    "200", "4,800"],
        ["TOTAL",  "800", "19,200"],
    ],
    col_widths=[1.5, 1.5, 1.5]
)
body(doc, "The rate limit status is shown in the UI as a progress bar per IOC type, along with time until reset.")

h2(doc, "3.7 Step-by-Step Enrichment Flow")
body(doc, "When the user clicks 'Enrich IOCs':")

steps = [
    ("Step 1", "Fetch Raw IOCs",          "Pull fresh IOCs from all 14 feeds"),
    ("Step 2", "Pre-Filter",              "Remove CIDR blocks, deduplicate same IOC values"),
    ("Step 3", "Rate Limit Check",        "Apply 200/type/hour limit, queue excess as 'rate-limited'"),
    ("Step 4", "Parallel Enrichment",     "Call all sources simultaneously using asyncio.gather()"),
    ("Step 5", "Score Calculation",       "BASE + all deltas, clamped to 0-100"),
    ("Step 6", "Confirmation Filter",     "Remove IOCs with 0 source confirmations (false positives)"),
    ("Step 7", "Sort & Return",           "Sort by risk_score descending, return to frontend"),
]
add_table(doc,
    ["Step", "Name", "Description"],
    steps,
    col_widths=[0.6, 1.8, 4.8]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — LIVE SEARCH
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Tab 3 — Live Search")
divider(doc)

body(doc,
    "The Live Search tab allows analysts to search for any specific IOC — an IP address, domain, "
    "file hash, or URL — and get results from both the local database and external threat intelligence."
)

h2(doc, "How Live Search Works")
bullet(doc, "User types an IOC value in the search box")
bullet(doc, "The system auto-detects the IOC type (IP, domain, hash, URL)")
bullet(doc, "Searches the local PostgreSQL database for any tracked IOC matching that value")
bullet(doc, "Simultaneously queries external sources (OTX, lookup APIs) for enriched context")
bullet(doc, "Returns merged results with source attribution")

h2(doc, "IOC Type Auto-Detection")
add_table(doc,
    ["Pattern", "Detected Type", "Example"],
    [
        ["Matches IPv4 regex",     "IP Address", "192.168.1.1"],
        ["Contains '/' with path", "URL",        "http://malware.com/payload.exe"],
        ["MD5: 32 hex chars",      "File Hash",  "d41d8cd98f00b204e9800998ecf8427e"],
        ["SHA256: 64 hex chars",   "File Hash",  "e3b0c44298fc1c149afb..."],
        ["Everything else",        "Domain",     "malware-c2.ru"],
    ],
    col_widths=[2.0, 1.5, 3.7]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — TRACKED INDICATORS
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Tab 4 — Tracked Indicators")
divider(doc)

body(doc,
    "Tracked Indicators are IOCs that your security team has manually added to the portal's "
    "local database. Unlike raw feed IOCs (which come from global feeds), tracked IOCs are "
    "specific to your organization's threat landscape."
)

h2(doc, "Why Track IOCs?")
bullet(doc, "IOC was found in your network logs or endpoint telemetry")
bullet(doc, "IOC was shared by a partner organization or ISAC (Information Sharing and Analysis Center)")
bullet(doc, "IOC is part of an active incident response investigation")
bullet(doc, "IOC targets your specific industry sector or brand")

h2(doc, "Tracked IOC Fields")
add_table(doc,
    ["Field", "Type", "Description"],
    [
        ["value",       "String",   "The IOC value (IP, domain, hash, or URL)"],
        ["ioc_type",    "Enum",     "Type: ip / domain / hash / url"],
        ["severity",    "Enum",     "critical / high / medium / low"],
        ["description", "Text",     "Analyst notes about this IOC"],
        ["source",      "String",   "Where this IOC came from"],
        ["created_at",  "DateTime", "When it was added to the database"],
        ["created_by",  "FK User",  "Which analyst added it"],
    ],
    col_widths=[1.3, 1.0, 4.9]
)

h2(doc, "Access Control")
body(doc, "Adding tracked IOCs requires Analyst or Admin role. Deletion requires Admin role.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — BACKEND ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Backend Architecture")
divider(doc)

h2(doc, "Key Files")
add_table(doc,
    ["File", "Purpose"],
    [
        ["backend/services/raw_ioc_feed.py", "Fetches all 14 raw IOC feeds concurrently"],
        ["backend/services/ioc_scorer.py",   "Enriches IOCs with risk scores from multiple sources"],
        ["backend/routes/admin.py",          "REST API endpoints for all IOC operations"],
        ["frontend/src/pages/IOCManagement.jsx", "Full frontend UI with all 4 tabs"],
    ],
    col_widths=[2.8, 4.4]
)

h2(doc, "Concurrency Model")
body(doc,
    "Both the raw feed fetcher and the enrichment engine use Python's asyncio for concurrent "
    "HTTP requests. This means all sources are queried simultaneously, not sequentially."
)
code_block(doc, "Sequential:  14 sources × 2s each = 28 seconds")
code_block(doc, "Concurrent:  14 sources at same time ≈ 2-4 seconds (fastest source determines total time)")
body(doc, "A Semaphore limits concurrency to prevent overwhelming APIs:")
code_block(doc, "asyncio.Semaphore(10)  →  max 10 IOCs enriched simultaneously")
code_block(doc, "asyncio.Semaphore(3)   →  ip-api.com specific throttle (45 req/min limit)")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — API ENDPOINTS
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "7. API Endpoints")
divider(doc)

add_table(doc,
    ["Method", "Endpoint", "Description", "Auth"],
    [
        ["GET",  "/api/admin/iocs",                      "List all tracked IOCs (paginated)",           "Required"],
        ["POST", "/api/admin/iocs",                      "Add a new tracked IOC",                       "Analyst+"],
        ["DELETE","/api/admin/iocs/{id}",               "Delete a tracked IOC",                        "Admin"],
        ["GET",  "/api/admin/iocs/raw-feed",             "Fetch live raw IOC feed from all 14 sources", "Required"],
        ["GET",  "/api/admin/iocs/enriched-feed",        "Fetch enriched IOCs with risk scores",        "Required"],
        ["GET",  "/api/admin/iocs/enrichment-rate-status","Get rate limit status per IOC type",          "Required"],
        ["GET",  "/api/admin/iocs/live-search",          "Search IOCs across local DB + external",      "Required"],
    ],
    col_widths=[0.8, 2.8, 2.5, 1.0]
)

h2(doc, "Query Parameters")
add_table(doc,
    ["Endpoint", "Parameter", "Values", "Description"],
    [
        ["/raw-feed",      "ioc_type", "ip, domain, hash, url",           "Filter by IOC type"],
        ["/raw-feed",      "severity", "critical, high, medium, low",      "Filter by severity"],
        ["/raw-feed",      "source",   "URLHaus, FeodoTracker, ...",       "Filter by specific feed"],
        ["/raw-feed",      "limit",    "integer (default 200, max 500)",   "IOCs per source"],
        ["/enriched-feed", "ioc_type", "ip, domain, hash, url",           "Enrich specific type only"],
        ["/enriched-feed", "limit",    "integer (default 100)",            "Max IOCs to enrich"],
        ["/admin/iocs",    "skip",     "integer (default 0)",              "Pagination offset"],
        ["/admin/iocs",    "limit",    "integer (default 50)",             "Page size"],
    ],
    col_widths=[1.6, 1.2, 2.0, 2.4]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8 — ENV VARIABLES
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Environment Variables & API Keys")
divider(doc)

body(doc, "The following API keys are needed in backend/.env for full enrichment functionality:")

add_table(doc,
    ["Variable", "Service", "Cost", "How to Get"],
    [
        ["ABUSEIPDB_API_KEY",       "AbuseIPDB",      "Free",    "Register at abuseipdb.com → API section"],
        ["ALIENVAULT_OTX_API_KEY",  "AlienVault OTX", "Free",    "Register at otx.alienvault.com → Settings → API Key"],
    ],
    col_widths=[2.0, 1.5, 0.8, 2.9]
)

body(doc, "Sources that require NO key (automatically work):")
bullet(doc, "Shodan InternetDB — https://internetdb.shodan.io/{ip}")
bullet(doc, "GreyNoise Community — https://api.greynoise.io/v3/community/{ip}")
bullet(doc, "ip-api.com — http://ip-api.com/json/{ip}")
bullet(doc, "MalwareBazaar — https://mb-api.abuse.ch/api/v1/")
bullet(doc, "All 14 raw feed sources (URLHaus, FeodoTracker, Spamhaus, etc.)")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 9 — LIMITATIONS
# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Limitations & Notes")
divider(doc)

h2(doc, "Known Limitations")
add_table(doc,
    ["Limitation", "Impact", "Mitigation"],
    [
        ["URLHaus /v1/host/ now requires auth (401)",     "Cannot do per-IOC URL/domain lookup via URLHaus",   "URLHaus still used for raw feed (bulk CSV)"],
        ["ThreatFox search_ioc requires auth (401)",      "Cannot search individual IOCs on ThreatFox",        "ThreatFox still used for raw feed"],
        ["Spamhaus DROP gives CIDR blocks, not plain IPs","Cannot enrich network ranges like 5.0.0.0/8",      "CIDR blocks automatically filtered and skipped"],
        ["ip-api.com limit: 45 req/min",                  "Rate-limited when enriching many IPs at once",      "Dedicated semaphore (max 3 concurrent calls)"],
        ["AbuseIPDB limit: 1,000/day",                    "May run out if enriching >1000 IPs in one day",    "Counted in rate limiter; others still work"],
        ["Enrichment not real-time cached",               "Each click re-fetches from external APIs",          "Consider adding Redis cache for production"],
        ["In-memory rate limiter",                        "Rate limit resets on server restart",               "Use Redis or DB for persistent rate tracking"],
    ],
    col_widths=[2.0, 2.2, 2.8]
)

h2(doc, "Recommendations for Production")
bullet(doc, "Add Redis caching for enrichment results (TTL: 1-4 hours) to reduce API calls")
bullet(doc, "Register for MetaDefender (OPSWAT) free API — 5,000 lookups/day, closest to VirusTotal")
bullet(doc, "Consider AlienVault OTX as primary enrichment for all IOC types (very generous limits)")
bullet(doc, "Set up daily automated enrichment job for all tracked IOCs in the database")
bullet(doc, "Add webhook alerts when a raw feed IOC matches a tracked indicator")

h2(doc, "Data Freshness")
add_table(doc,
    ["Data Type", "Freshness", "Notes"],
    [
        ["Raw IOC Feed",       "Real-time (on-demand)",  "Fetched fresh every time the tab loads"],
        ["Enriched Results",   "Real-time (on-demand)",  "Fetched fresh when 'Enrich IOCs' is clicked"],
        ["Tracked IOCs",       "Persistent (database)",  "Stored in PostgreSQL, never expires"],
        ["Rate Limit Counter", "In-memory, resets hourly","Lost on server restart (by design)"],
    ],
    col_widths=[1.8, 1.8, 3.6]
)

# ─── Footer ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
divider(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(
    f"SecureEye Portal — IOC Management Documentation  |  "
    f"Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Confidential"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = "IOC_Management_Documentation.docx"
doc.save(output_path)
print(f"[OK] Documentation saved to: {output_path}")
